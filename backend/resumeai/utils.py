from io import BytesIO

from docx import Document
from docx.shared import Pt
from pypdf import PdfReader


def _sanitize_text(text: str) -> str:
    """Remove NUL and other control chars that break DB/JSON."""
    if not text:
        return text
    return text.replace("\x00", "").replace("\r\n", "\n").replace("\r", "\n")


def extract_text_from_docx(path_or_file) -> str:
    """Accept file path (str) or file-like object."""
    doc = Document(path_or_file)
    parts: list[str] = []
    for paragraph in doc.paragraphs:
        text = _sanitize_text((paragraph.text or "").strip())
        if text:
            parts.append(text)
    return _sanitize_text("\n".join(parts))


def extract_text_from_pdf(path_or_file) -> str:
    """Accept file path (str) or file-like object."""
    reader = PdfReader(path_or_file)
    parts: list[str] = []
    for page in reader.pages:
        text = _sanitize_text((page.extract_text() or "").strip())
        if text:
            parts.append(text)
    return _sanitize_text("\n".join(parts))


def create_docx_from_text(text: str) -> bytes:
    """Create a DOCX file from plain text. Each line becomes a paragraph."""
    doc = Document()
    style = doc.styles["Normal"]
    style.font.size = Pt(11)
    for line in (text or "").splitlines():
        doc.add_paragraph(line.strip() or "")
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()


def pdf_to_html(path_or_file) -> str:
    """Convert PDF to editable HTML. Accept file path (str) or file-like object (read as bytes)."""
    import pymupdf
    if hasattr(path_or_file, "read"):
        data = path_or_file.read()
        doc = pymupdf.open(stream=data, filetype="pdf")
    else:
        doc = pymupdf.open(path_or_file)
    try:
        parts = []
        for page in doc:
            html = page.get_text("html")
            if html:
                parts.append(html)
        return _sanitize_text("\n".join(parts)) if parts else ""
    finally:
        doc.close()


def html_to_pdf(html: str) -> bytes:
    """Convert HTML to PDF. Returns PDF bytes."""
    from xhtml2pdf import pisa
    result = BytesIO()
    style = """
    body { font-family: Helvetica, Arial, sans-serif; font-size: 11pt; line-height: 1.4; margin: 1in; }
    p { margin: 0.3em 0; }
    ul, ol { margin: 0.3em 0; padding-left: 1.5em; }
    strong { font-weight: bold; }
    em { font-style: italic; }
    """
    html_src = f"""<!DOCTYPE html><html><head><meta charset="utf-8"/><style>{style}</style></head><body>{html or ""}</body></html>"""
    pdf = pisa.CreatePDF(html_src, dest=result, encoding="utf-8")
    if pdf.err:
        raise ValueError("PDF creation failed")
    result.seek(0)
    return result.getvalue()